import streamlit as st
import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io
from datetime import datetime
import easyocr

# --- 1. CONFIGURACIÓN Y ESTILO ---
st.set_page_config(page_title="Asistente Prof. Cárdenas", layout="wide")

st.markdown(
    """
    <style>
    .stApp { background-color: #E3F2FD; }
    .foto-perfil { position: fixed; top: 50px; right: 30px; z-index: 1000; }
    .foto-perfil img { width: 115px; height: 115px; border-radius: 50%; border: 3px solid #1976D2; object-fit: cover; }
    </style>
    <div class="foto-perfil">
        <img src="https://raw.githubusercontent.com/matematicocardenas25-cripto/Asistente-de-formatos-educativos-/main/foto.jpg.jpeg">
    </div>
    """, unsafe_allow_html=True
)

# --- 2. LÓGICA DE CONTENIDO AUTOMÁTICO ---
def generar_conclusiones(unidad):
    return (f"Se desarrolló el contenido de {unidad}. Se dio a conocer los fundamentos teóricos para llevarlos "
            "posteriormente a la práctica. Con esta sesión el estudiante está en capacidad de utilizar los "
            "diferentes instrumentos y métodos de recolección de datos.")

def generar_recomendaciones(unidad):
    return (f"Revisar bibliografía para profundizar conocimiento del tema impartido ({unidad}) durante la sesión de clase.")

# --- 3. GENERADOR DE LATEX ---
def generar_latex(d):
    concl = generar_conclusiones(d['unidad'])
    recom = generar_recomendaciones(d['unidad'])
    
    latex_code = f"""\\documentclass[12pt,a4paper]{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage[spanish]{{babel}}
\\usepackage{{geometry}}
\\geometry{{left=2.5cm,right=2.5cm,top=2.5cm,bottom=2.5cm}}

\\begin{{document}}

\\begin{{center}}
    \\textbf{{PROGRAMACIÓN DIDÁCTICA PARA LOS APRENDIZAJES}}
\\end{{center}}

\\section*{{I. DATOS GENERALES:}}
\\textbf{{1.1 Área de conocimiento:}} {d['area']} \\\\
\\textbf{{1.2 Carrera:}} {d['carrera']} \\hfill \\textbf{{1.3 Modalidad:}} {d['modalidad']} \\quad \\textbf{{Turno:}} Diurno \\\\
\\textbf{{1.4. Nombre de la asignatura:}} {d['asignatura']} \\\\
\\textbf{{1.5. Fecha:}} {d['fecha']} \\hfill \\textbf{{1.6. Hora:}} 10:30 am – 1:00 pm \\\\
\\textbf{{1.7. Profesor (a):}} {d['profesor']}

\\section*{{II. UNIDAD:}}
{d['unidad']}

\\section*{{III. OBJETIVO GENERAL:}}
Reconocer diferentes métodos para la recolección y organización de información para la construcción de base de datos mediante técnicas descripticas.

\\section*{{IV. OBJETIVO(S) ESPECÍFICO(S):}}
\\begin{{itemize}}
    \\item Definir conceptos básicos de estadística, fuentes de datos para investigación.
    \\item Explicar conceptos básicos de estadística, fuentes de datos para investigación, necesidad para realizar una investigación.
    \\item Aplicar conceptos básicos de estadística y obtención de datos mediante encuesta.
\\end{{itemize}}

\\section*{{V. EVALUACIÓN DE LOS APRENDIZAJES (Criterios y Evidencias):}}
{d['evaluacion']}

\\section*{{VI. ACTIVIDADES DEL DOCENTE Y DE LOS ESTUDIANTES:}}
{d['actividades']}

\\section*{{VII. MEDIOS O RECURSOS DIDÁCTICOS NECESARIOS:}}
Plan de clase, Plan calendario, Programa de asignatura, Libro en físico y digital, borrador, lapiceros, pizarra acrílica, entre otros.

\\section*{{VIII. CONCLUSIONES}}
{concl}

\\section*{{IX. RECOMENDACIONES:}}
{recom}

\\section*{{X. BIBLIOGRAFIA:}}
{d['biblio']}

\\end{{document}}
"""
    return latex_code

# --- 4. GENERADOR DE WORD ---
def generar_word_oficial(d):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.add_run("PROGRAMACIÓN DIDÁCTICA PARA LOS APRENDIZAJES").bold = True

    doc.add_heading('I. DATOS GENERALES:', level=1)
    p1 = doc.add_paragraph()
    p1.add_run("1.1 Área de conocimiento: ").bold = True
    p1.add_run(f"{d['area']}")

    p2 = doc.add_paragraph()
    p2.add_run("1.2 Carrera: ").bold = True
    p2.add_run(f"{d['carrera']}")
    p2.add_run("    1.3 Modalidad: ").bold = True
    p2.add_run(f"{d['modalidad']}")
    p2.add_run("    Turno: ").bold = True
    p2.add_run("Diurno")

    p3 = doc.add_paragraph()
    p3.add_run("1.4. Nombre de la asignatura: ").bold = True
    p3.add_run(f"{d['asignatura']}")

    p4 = doc.add_paragraph()
    p4.add_run("1.5. Fecha: ").bold = True
    p4.add_run(f"{d['fecha']}")
    p4.add_run("    1.6. Hora: ").bold = True
    p4.add_run("10:30 am – 1:00 pm")

    p5 = doc.add_paragraph()
    p5.add_run("1.7. Profesor (a): ").bold = True
    p5.add_run(f"{d['profesor']}")

    doc.add_heading('II. UNIDAD:', level=1)
    doc.add_paragraph(d['unidad'])

    doc.add_heading('III. OBJETIVO GENERAL:', level=1)
    doc.add_paragraph("Reconocer diferentes métodos para la recolección y organización de información para la construcción de base de datos mediante técnicas descripticas.")

    doc.add_heading('IV. OBJETIVO(S) ESPECÍFICO(S):', level=1)
    for obj in ["Definir conceptos básicos.", "Explicar fuentes de datos.", "Aplicar encuestas."]:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_heading('V. EVALUACIÓN DE LOS APRENDIZAJES (Criterios y Evidencias):', level=1)
    doc.add_paragraph(d['evaluacion'])

    doc.add_heading('VI. ACTIVIDADES DEL DOCENTE Y DE LOS ESTUDIANTES:', level=1)
    doc.add_paragraph(d['actividades'])

    doc.add_heading('VII. MEDIOS O RECURSOS DIDÁCTICOS NECESARIOS:', level=1)
    doc.add_paragraph("Plan de clase, Plan calendario, Programa de asignatura, Libro en físico y digital, borrador, lapiceros, pizarra acrílica, entre otros.")

    doc.add_heading('VIII. CONCLUSIONES', level=1)
    doc.add_paragraph(generar_conclusiones(d['unidad']))

    doc.add_heading('IX. RECOMENDACIONES:', level=1)
    doc.add_paragraph(generar_recomendaciones(d['unidad']))

    doc.add_heading('X. BIBLIOGRAFIA:', level=1)
    doc.add_paragraph(d['biblio'])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# --- 5. INTERFAZ ---
tab1, tab2 = st.tabs(["📄 Planificación Didáctica", "📊 Calculadora Multidimensión"])

with tab1:
    st.title("Generador de Formatos Prof. Cárdenas")
    archivo_img = st.file_uploader("📷 Escanear Actividades", type=['jpg','png','jpeg'])
    texto_escaneado = ""
    if archivo_img:
        reader = easyocr.Reader(['es'])
        texto_escaneado = "\n".join(reader.readtext(np.array(Image.open(archivo_img)), detail=0))

    with st.form("form_doc"):
        c1, c2 = st.columns(2)
        area = c1.text_input("1.1 Área de conocimiento", "Ciencias Económica y Empresariales, Ingeniería y Construcción")
        carrera = c2.text_input("1.2 Carrera", "Todas")
        asignatura = c1.text_input("1.4 Nombre de la asignatura", "Estadística descriptiva")
        profesor = c2.text_input("1.7 Profesor (a)", "Ismael Antonio Cárdenas López")
        fecha = st.text_input("1.5 Fecha", datetime.now().strftime("%d/%m/%Y"))
        modalidad = "Presencial"
        unidad = st.text_area("II. UNIDAD", "Recopilación de datos. Clase practica 1.")
        evaluacion = st.text_area("V. EVALUACIÓN", "Identifique los diferentes tipos de variables. Registro de participación.")
        actividades = st.text_area("VI. ACTIVIDADES", value=texto_escaneado, height=150)
        biblio = st.text_area("X. BIBLIOGRAFIA", "Saldaña, M. Y. (2024). Principios de Estadística descriptiva. Perú.")
        
        btn_validar = st.form_submit_button("✅ Procesar Documentos")

    if btn_validar:
        datos = {
            "area": area, "carrera": carrera, "asignatura": asignatura, "profesor": profesor,
            "fecha": fecha, "modalidad": modalidad, "unidad": unidad, "evaluacion": evaluacion,
            "actividades": actividades, "biblio": biblio
        }
        st.success("¡Formatos listos para descargar!")
        col1, col2 = st.columns(2)
        col1.download_button("📥 Descargar Word (.docx)", generar_word_oficial(datos), f"Plan_{asignatura}.docx")
        col2.download_button("📥 Descargar LaTeX (.tex)", generar_latex(datos).encode('utf-8'), f"Plan_{asignatura}.tex")

with tab2:
    # --- CALCULADORA (SIN CAMBIOS) ---
    st.title("📊 Gráficos Trascendentales (Ejes 0,0)")
    dim = st.radio("Dimensión:", ["2D (Plano)", "3D (Espacial)"], horizontal=True)
    contexto_mat = {"np": np, "x": None, "y": None, "sin": np.sin, "cos": np.cos, "tan": np.tan, "exp": np.exp, "log": np.log, "sqrt": np.sqrt}
    col_a, col_b = st.columns([1, 2])
    with col_a:
        if dim == "2D (Plano)":
            eq = st.text_input("f(x) =", "sin(x)")
            r = st.slider("Rango", -50, 50, (-10, 10))
        else:
            eq_3d = st.text_input("z = f(x, y)", "x**2 - y**2")
            res = st.slider("Resolución", 5, 20, 10)
    with col_b:
        try:
            if dim == "2D (Plano)":
                x_val = np.linspace(r[0], r[1], 500)
                contexto_mat["x"] = x_val
                y_val = eval(eq, {"__builtins__": None}, contexto_mat)
                fig = go.Figure(go.Scatter(x=x_val, y=y_val, line=dict(color='#1976D2', width=3)))
                fig.update_xaxes(zeroline=True, zerolinewidth=2, zerolinecolor='black')
                fig.update_yaxes(zeroline=True, zerolinewidth=2, zerolinecolor='black')
                st.plotly_chart(fig)
            else:
                x = y = np.linspace(-res, res, 100)
                X, Y = np.meshgrid(x, y)
                contexto_mat["x"], contexto_mat["y"] = X, Y
                Z = eval(eq_3d, {"__builtins__": None}, contexto_mat)
                st.plotly_chart(go.Figure(data=[go.Surface(z=Z, x=X, y=Y)]))
        except Exception as e:
            st.error(f"Error matemático: {e}")
